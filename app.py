import streamlit as st
import PyPDF2
import docx
import json
import sqlite3
import hashlib
import uuid
from datetime import datetime
from typing import List, Dict, Any
import openai
from langchain.text_splitter import RecursiveCharacterTextSplitter
import plotly.express as px
import plotly.graph_objects as go
import pandas as pd
import re
import time
import random
import io

# Configure Streamlit page
st.set_page_config(
    page_title="🧠 Sensai AI - Doc→Quiz Bot",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize OpenAI with hackathon configuration
if 'openai_api_key' not in st.session_state:
    st.session_state.openai_api_key = st.secrets.get("OPENAI_API_KEY", "")

# Initialize OpenAI client with hackathon settings
def initialize_openai_client(api_key):
    """Initialize OpenAI client with hackathon configuration"""
    try:
        return openai.OpenAI(
            api_key=api_key,
            base_url="https://agent.dev.hyperverge.org"
        )
    except Exception as e:
        st.error(f"Failed to initialize OpenAI client: {str(e)}")
        return None

# Database setup for sensai features
def init_sensai_db():
    conn = sqlite3.connect('sensai_quiz.db')
    c = conn.cursor()
    
    # Quiz attempts table
    c.execute('''CREATE TABLE IF NOT EXISTS quiz_attempts
                 (id INTEGER PRIMARY KEY, 
                  session_id TEXT,
                  user_id TEXT,
                  file_hash TEXT,
                  filename TEXT,
                  difficulty TEXT,
                  score INTEGER,
                  total_questions INTEGER,
                  timestamp TEXT,
                  chunks_covered TEXT,
                  wrong_questions TEXT,
                  time_taken REAL)''')
    
    # Question items table (item bank)
    c.execute('''CREATE TABLE IF NOT EXISTS question_items
                 (id INTEGER PRIMARY KEY,
                  file_hash TEXT,
                  question_text TEXT,
                  option_a TEXT,
                  option_b TEXT,
                  option_c TEXT,
                  option_d TEXT,
                  correct_answer TEXT,
                  explanation TEXT,
                  source_chunk TEXT,
                  page_number INTEGER,
                  chunk_id INTEGER,
                  difficulty TEXT,
                  citation TEXT)''')
    
    # Chat sessions table
    c.execute('''CREATE TABLE IF NOT EXISTS chat_sessions
                 (id INTEGER PRIMARY KEY,
                  session_id TEXT,
                  user_id TEXT,
                  current_question INTEGER,
                  score INTEGER,
                  questions_order TEXT,
                  wrong_questions TEXT,
                  started_at TEXT,
                  completed_at TEXT)''')
    
    # Debug logs table
    c.execute('''CREATE TABLE IF NOT EXISTS debug_logs
                 (id INTEGER PRIMARY KEY,
                  timestamp TEXT,
                  level TEXT,
                  message TEXT,
                  data TEXT)''')
    
    conn.commit()
    conn.close()

init_sensai_db()

# Debug logging function
def log_debug(level: str, message: str, data: Any = None):
    """Log debug information to database and display"""
    conn = sqlite3.connect('sensai_quiz.db')
    c = conn.cursor()
    
    c.execute('''INSERT INTO debug_logs (timestamp, level, message, data)
                 VALUES (?, ?, ?, ?)''',
              (datetime.now().isoformat(), level, message, 
               json.dumps(data, default=str) if data else None))
    conn.commit()
    conn.close()
    
    # Also display in Streamlit based on level
    if level == "ERROR":
        st.error(f"🐛 DEBUG: {message}")
    elif level == "WARNING":
        st.warning(f"⚠️ DEBUG: {message}")
    elif level == "INFO":
        st.info(f"ℹ️ DEBUG: {message}")

# SENSAI GPT Prompt Templates (Improved)
SENSAI_PROMPTS = {
    "easy": """You are an expert quiz generator. Based on this text, create exactly 2 multiple choice questions for basic understanding.

TEXT TO ANALYZE:
{chunk}

REQUIREMENTS:
- Create exactly 2 questions
- Focus on basic facts, definitions, and direct information
- Make questions clear and unambiguous
- Provide 4 distinct options (A, B, C, D)
- One correct answer per question
- Include brief explanations

FORMAT EXACTLY AS SHOWN:
QUESTION: What is the main topic discussed in this section?
A) Option 1
B) Option 2  
C) Option 3
D) Option 4
CORRECT: A
EXPLANATION: This is correct because the text clearly states...
CITATION: Page {page_num}: "Brief relevant quote"
---
QUESTION: [Second question following same format]
A) Option 1
B) Option 2
C) Option 3  
D) Option 4
CORRECT: B
EXPLANATION: This is correct because...
CITATION: Page {page_num}: "Brief relevant quote"
---""",

    "medium": """You are an expert quiz generator. Based on this text, create exactly 2 multiple choice questions for comprehension and application.

TEXT TO ANALYZE:
{chunk}

REQUIREMENTS:
- Create exactly 2 questions
- Focus on understanding concepts and relationships
- Test comprehension, not just recall
- Make questions require understanding, not just memorization
- Provide 4 distinct options (A, B, C, D)
- Include detailed explanations

FORMAT EXACTLY AS SHOWN:
QUESTION: How does concept X relate to concept Y based on the text?
A) Option 1
B) Option 2
C) Option 3
D) Option 4
CORRECT: C
EXPLANATION: This is correct because the text explains the relationship...
CITATION: Page {page_num}: "Brief relevant quote"
---
QUESTION: [Second question following same format]
A) Option 1
B) Option 2
C) Option 3
D) Option 4
CORRECT: A
EXPLANATION: This demonstrates understanding because...
CITATION: Page {page_num}: "Brief relevant quote"
---""",

    "hard": """You are an expert quiz generator. Based on this text, create exactly 2 multiple choice questions for analysis and critical thinking.

TEXT TO ANALYZE:
{chunk}

REQUIREMENTS:
- Create exactly 2 questions
- Focus on analysis, synthesis, and evaluation
- Test ability to connect concepts and think critically
- Questions should require deep understanding
- Provide 4 distinct options (A, B, C, D)
- Include comprehensive explanations

FORMAT EXACTLY AS SHOWN:
QUESTION: What can be inferred about [complex concept] based on the evidence presented?
A) Option 1
B) Option 2
C) Option 3
D) Option 4
CORRECT: D
EXPLANATION: This requires critical analysis because...
CITATION: Page {page_num}: "Brief relevant quote"
---
QUESTION: [Second question following same format]
A) Option 1
B) Option 2
C) Option 3
D) Option 4
CORRECT: B
EXPLANATION: This demonstrates synthesis of ideas because...
CITATION: Page {page_num}: "Brief relevant quote"
---"""
}

class SensaiDocumentProcessor:
    @staticmethod
    def extract_text_from_pdf(file) -> tuple[str, int, List[Dict]]:
        """Extract text from PDF with page tracking for citations using PyPDF2"""
        try:
            # Reset file pointer
            file.seek(0)
            
            # Use PyPDF2 instead of PyMuPDF for better Streamlit compatibility
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
            text = ""
            page_count = len(pdf_reader.pages)
            page_content = []
            
            log_debug("INFO", f"PDF detected: {page_count} pages")
            
            # Limit to first 10 pages to prevent API overuse
            max_pages = min(10, page_count)
            
            if page_count > 10:
                st.warning(f"🔍 Large doc detected ({page_count} pages). Processing first 10 pages only.")
            
            for page_num in range(max_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    # Clean extracted text
                    page_text = page_text.replace('\n\n', '\n').strip()
                    
                    if page_text and len(page_text) > 50:  # Only add if page has substantial text
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text + "\n"
                        
                        page_content.append({
                            "page_num": page_num + 1,
                            "text": page_text,
                            "word_count": len(page_text.split())
                        })
                        
                        log_debug("INFO", f"Page {page_num + 1}: {len(page_text)} chars, {len(page_text.split())} words")
                except Exception as page_error:
                    log_debug("ERROR", f"Could not process page {page_num + 1}", str(page_error))
                    continue
            
            log_debug("INFO", f"Total extracted text: {len(text)} characters")
            return text, page_count, page_content
            
        except Exception as e:
            log_debug("ERROR", "PDF processing failed", str(e))
            return "", 0, []

    @staticmethod
    def extract_text_from_docx(file) -> tuple[str, int, List[Dict]]:
        """Extract text from DOCX with page estimation"""
        try:
            doc = docx.Document(file)
            text = ""
            para_count = 0
            page_content = []
            current_page = 1
            chars_per_page = 2500  # Rough estimation
            current_chars = 0
            current_page_text = ""
            
            log_debug("INFO", f"DOCX detected: {len(doc.paragraphs)} paragraphs")
            
            for para in doc.paragraphs:
                para_text = para.text + "\n"
                if len(para_text.strip()) > 0:  # Only add non-empty paragraphs
                    text += para_text
                    current_page_text += para_text
                    current_chars += len(para_text)
                    para_count += 1
                
                # Estimate page break
                if current_chars > chars_per_page:
                    if current_page_text.strip():
                        page_content.append({
                            "page_num": current_page,
                            "text": current_page_text,
                            "word_count": len(current_page_text.split())
                        })
                        log_debug("INFO", f"DOCX Page {current_page}: {len(current_page_text)} chars")
                    
                    current_page += 1
                    current_chars = 0
                    current_page_text = ""
                    
                    # Limit to 10 pages
                    if current_page > 10:
                        break
            
            # Add final page if exists
            if current_page_text.strip():
                page_content.append({
                    "page_num": current_page,
                    "text": current_page_text,
                    "word_count": len(current_page_text.split())
                })
            
            estimated_pages = len(page_content)
            log_debug("INFO", f"DOCX processing complete: {estimated_pages} pages, {len(text)} total chars")
            
            return text, estimated_pages, page_content
            
        except Exception as e:
            log_debug("ERROR", "DOCX processing failed", str(e))
            return "", 0, []

    @staticmethod
    def chunk_text_with_citations(text: str, page_content: List[Dict]) -> List[Dict[str, Any]]:
        """Split text into chunks with page number tracking for citations"""
        try:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=4000,  # Reduced chunk size for better processing
                chunk_overlap=200,
                length_function=len,
                separators=["\n\n", "\n", ". ", " ", ""]
            )
            
            chunks = splitter.split_text(text)
            log_debug("INFO", f"Text split into {len(chunks)} chunks")
            
            # Create chunk objects with page mapping
            chunk_objects = []
            for i, chunk in enumerate(chunks):
                if len(chunk.strip()) < 100:  # Skip very small chunks
                    continue
                    
                # Find which page this chunk primarily belongs to
                page_num = 1
                max_overlap = 0
                
                for page in page_content:
                    # Calculate overlap between chunk and page
                    chunk_words = set(chunk.lower().split())
                    page_words = set(page["text"].lower().split())
                    overlap = len(chunk_words.intersection(page_words))
                    
                    if overlap > max_overlap:
                        max_overlap = overlap
                        page_num = page["page_num"]
                
                chunk_obj = {
                    "id": i,
                    "text": chunk,
                    "page_num": page_num,
                    "word_count": len(chunk.split()),
                    "questions_generated": 0
                }
                
                chunk_objects.append(chunk_obj)
                log_debug("INFO", f"Chunk {i}: {len(chunk)} chars, page {page_num}")
            
            return chunk_objects
            
        except Exception as e:
            log_debug("ERROR", "Text chunking failed", str(e))
            return []

class SensaiQuizGenerator:
    @staticmethod
    def generate_mcqs_with_citations(chunk: Dict, difficulty: str = "medium") -> List[Dict]:
        """Generate MCQs with proper citations using Sensai prompts"""
        if not st.session_state.openai_api_key:
            log_debug("ERROR", "OpenAI API key not found")
            return []
        
        # Create snippet for citation
        snippet = chunk['text'][:100] + "..." if len(chunk['text']) > 100 else chunk['text']
        
        prompt = SENSAI_PROMPTS[difficulty].format(
            chunk=chunk['text'][:2000],  # Limit chunk size to prevent token overflow
            page_num=chunk['page_num'],
            snippet=snippet
        )
        
        try:
            # Initialize OpenAI client with hackathon configuration
            client = initialize_openai_client(st.session_state.openai_api_key)
            
            if not client:
                log_debug("ERROR", "Failed to initialize OpenAI client")
                return []
            
            log_debug("INFO", f"Generating questions for chunk {chunk['id']}, difficulty: {difficulty}")
            
            response = client.chat.completions.create(
                model="openai/gpt-4o-mini",  # Updated model name for hackathon
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1000,
                temperature=0.2  # Lower temperature for more consistent formatting
            )
            
            questions_text = response.choices[0].message.content
            log_debug("INFO", f"API Response received: {len(questions_text)} characters")
            
            # Parse and validate questions
            questions = SensaiQuizGenerator.parse_questions_with_citations(
                questions_text, chunk, difficulty
            )
            
            log_debug("INFO", f"Successfully parsed {len(questions)} questions from chunk {chunk['id']}")
            return questions
            
        except Exception as e:
            log_debug("ERROR", f"Question generation failed for chunk {chunk['id']}", str(e))
            return []
    
    @staticmethod
    def parse_questions_with_citations(questions_text: str, source_chunk: Dict, difficulty: str) -> List[Dict]:
        """Parse GPT response into structured questions with citations"""
        questions = []
        
        try:
            # Log the raw response for debugging
            log_debug("INFO", f"Parsing response: {questions_text[:200]}...")
            
            # Split by --- separator
            question_blocks = [block.strip() for block in questions_text.split("---") if block.strip()]
            
            log_debug("INFO", f"Found {len(question_blocks)} question blocks")
            
            for block_idx, block in enumerate(question_blocks):
                if not block.strip():
                    continue
                
                lines = [line.strip() for line in block.strip().split('\n') if line.strip()]
                
                if len(lines) < 6:  # Need minimum fields (question, A, B, C, D, correct)
                    log_debug("WARNING", f"Block {block_idx} too short: {len(lines)} lines")
                    continue
                
                question_data = {
                    "difficulty": difficulty, 
                    "chunk_id": source_chunk["id"],
                    "page_num": source_chunk["page_num"],
                    "source_chunk": source_chunk['text'][:300] + "...",
                    "id": len(questions)
                }
                
                # Parse each line
                for line in lines:
                    line = line.strip()
                    
                    if line.startswith("QUESTION:"):
                        question_data["question"] = line[9:].strip()
                    elif line.startswith("A)"):
                        question_data["A"] = line[2:].strip()
                    elif line.startswith("B)"):
                        question_data["B"] = line[2:].strip()
                    elif line.startswith("C)"):
                        question_data["C"] = line[2:].strip()
                    elif line.startswith("D)"):
                        question_data["D"] = line[2:].strip()
                    elif line.startswith("CORRECT:"):
                        question_data["correct"] = line[8:].strip()
                    elif line.startswith("EXPLANATION:"):
                        question_data["explanation"] = line[12:].strip()
                    elif line.startswith("CITATION:"):
                        question_data["citation"] = line[9:].strip()
                
                # Validate required fields
                required_fields = ["question", "A", "B", "C", "D", "correct"]
                missing_fields = [field for field in required_fields if field not in question_data or not question_data[field]]
                
                if missing_fields:
                    log_debug("WARNING", f"Block {block_idx} missing fields: {missing_fields}")
                    continue
                
                # Validate correct answer format
                if question_data["correct"] not in ["A", "B", "C", "D"]:
                    log_debug("WARNING", f"Block {block_idx} invalid correct answer: {question_data['correct']}")
                    continue
                
                # Set default values for optional fields
                if "explanation" not in question_data:
                    question_data["explanation"] = "No explanation provided."
                if "citation" not in question_data:
                    question_data["citation"] = f"Page {source_chunk['page_num']}: Source material"
                
                questions.append(question_data)
                log_debug("INFO", f"Successfully parsed question {len(questions)}: {question_data['question'][:50]}...")
            
            return questions
            
        except Exception as e:
            log_debug("ERROR", "Question parsing failed", str(e))
            return []

class SensaiQuizEngine:
    @staticmethod
    def create_quiz_session(questions: List[Dict], difficulty: str, user_id: str = "anonymous"):
        """Create a new quiz session"""
        session_id = str(uuid.uuid4())
        
        try:
            # Save session to database
            conn = sqlite3.connect('sensai_quiz.db')
            c = conn.cursor()
            
            # Randomize question order for adaptive experience
            question_order = list(range(len(questions)))
            random.shuffle(question_order)
            
            c.execute('''INSERT INTO chat_sessions 
                         (session_id, user_id, current_question, score, questions_order, 
                          wrong_questions, started_at)
                         VALUES (?, ?, ?, ?, ?, ?, ?)''',
                      (session_id, user_id, 0, 0, json.dumps(question_order), 
                       json.dumps([]), datetime.now().isoformat()))
            conn.commit()
            conn.close()
            
            log_debug("INFO", f"Quiz session created: {session_id}")
            return session_id, question_order
            
        except Exception as e:
            log_debug("ERROR", "Failed to create quiz session", str(e))
            return None, []

    @staticmethod
    def get_session_data(session_id: str):
        """Get session data from database"""
        try:
            conn = sqlite3.connect('sensai_quiz.db')
            c = conn.cursor()
            
            c.execute('SELECT * FROM chat_sessions WHERE session_id = ?', (session_id,))
            session = c.fetchone()
            conn.close()
            
            if session:
                return {
                    "session_id": session[1],
                    "user_id": session[2],
                    "current_question": session[3],
                    "score": session[4],
                    "questions_order": json.loads(session[5]),
                    "wrong_questions": json.loads(session[6]),
                    "started_at": session[7],
                    "completed_at": session[8]
                }
            return None
            
        except Exception as e:
            log_debug("ERROR", "Failed to get session data", str(e))
            return None

    @staticmethod
    def update_session(session_id: str, current_question: int, score: int, wrong_questions: List):
        """Update session progress"""
        try:
            conn = sqlite3.connect('sensai_quiz.db')
            c = conn.cursor()
            
            c.execute('''UPDATE chat_sessions 
                         SET current_question = ?, score = ?, wrong_questions = ?
                         WHERE session_id = ?''',
                      (current_question, score, json.dumps(wrong_questions), session_id))
            conn.commit()
            conn.close()
            
        except Exception as e:
            log_debug("ERROR", "Failed to update session", str(e))

class SensaiChatBot:
    @staticmethod
    def display_question_conversational(question: Dict, question_num: int, total: int):
        """Display question in conversational chat format"""
        
        # Bot message
        with st.chat_message("assistant", avatar="🤖"):
            st.markdown(f"**Question {question_num}/{total}** 📝")
            st.markdown(f"*Difficulty: {question.get('difficulty', 'medium').title()}*")
            st.markdown(f"**{question['question']}**")
            
            # Show citation
            if question.get('citation'):
                st.caption(f"📖 {question['citation']}")
        
        # User input area
        st.markdown("**Choose your answer:**")
        
        # Create columns for options
        col1, col2 = st.columns(2)
        
        selected_answer = None
        
        with col1:
            if st.button(f"🅰️ {question['A']}", key=f"A_{question_num}", use_container_width=True):
                selected_answer = 'A'
            if st.button(f"🅲 {question['C']}", key=f"C_{question_num}", use_container_width=True):
                selected_answer = 'C'
        
        with col2:
            if st.button(f"🅱️ {question['B']}", key=f"B_{question_num}", use_container_width=True):
                selected_answer = 'B'
            if st.button(f"🅳 {question['D']}", key=f"D_{question_num}", use_container_width=True):
                selected_answer = 'D'
        
        return selected_answer

    @staticmethod
    def show_conversational_feedback(is_correct: bool, question: Dict, selected_answer: str):
        """Show feedback in conversational format"""
        
        if is_correct:
            with st.chat_message("assistant", avatar="🎉"):
                messages = [
                    "Excellent! You got it right! 🎯",
                    "Perfect! That's the correct answer! ✨",
                    "Great job! You nailed it! 🌟",
                    "Awesome! Correct answer! 🚀"
                ]
                st.success(random.choice(messages))
                
                # Show explanation
                st.info(f"**Why it's correct:** {question.get('explanation', 'No explanation available.')}")
        else:
            with st.chat_message("assistant", avatar="🤔"):
                st.error("Oops! That's not quite right. 😅")
                st.info(f"You selected: **{selected_answer}) {question[selected_answer]}**")
                st.info(f"Correct answer: **{question['correct']}) {question[question['correct']]}**")
                
                # Show explanation immediately
                with st.expander("🔍 Explanation & Source", expanded=True):
                    st.markdown(f"**Explanation:** {question.get('explanation', 'No explanation available.')}")
                    if question.get('citation'):
                        st.markdown(f"**Source:** {question['citation']}")
                    st.markdown(f"**Context:** {question.get('source_chunk', 'No additional context available.')}")

class SensaiVisualizer:
    @staticmethod
    def show_coverage_heatmap(chunks: List[Dict], questions: List[Dict]):
        """Enhanced coverage heatmap visualization"""
        
        # Count questions per chunk/page
        chunk_stats = {}
        page_stats = {}
        
        for question in questions:
            chunk_id = question.get('chunk_id', 0)
            page_num = question.get('page_num', 1)
            
            chunk_stats[chunk_id] = chunk_stats.get(chunk_id, 0) + 1
            page_stats[page_num] = page_stats.get(page_num, 0) + 1
        
        # Create visualizations
        col1, col2 = st.columns(2)
        
        with col1:
            # Chunk coverage bar chart
            if chunks:
                chunk_data = pd.DataFrame({
                    'Chunk': list(range(len(chunks))),
                    'Questions': [chunk_stats.get(i, 0) for i in range(len(chunks))]
                })
                
                fig_chunks = px.bar(
                    chunk_data, 
                    x='Chunk', 
                    y='Questions',
                    title="📊 Questions per Chunk",
                    color='Questions',
                    color_continuous_scale='Viridis'
                )
                fig_chunks.update_layout(height=400)
                st.plotly_chart(fig_chunks, use_container_width=True)
        
        with col2:
            # Page coverage heatmap
            if page_stats:
                pages = sorted(page_stats.keys())
                page_data = pd.DataFrame({
                    'Page': pages,
                    'Questions': [page_stats[p] for p in pages]
                })
                
                fig_pages = px.bar(
                    page_data,
                    x='Page',
                    y='Questions', 
                    title="📄 Questions per Page",
                    color='Questions',
                    color_continuous_scale='Blues'
                )
                fig_pages.update_layout(height=400)
                st.plotly_chart(fig_pages, use_container_width=True)
        
        # Coverage statistics
        st.subheader("📈 Coverage Insights")
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            total_chunks = len(chunks) if chunks else 0
            st.metric("Total Chunks", total_chunks)
        
        with col2:
            covered_chunks = len([c for c in chunk_stats.values() if c > 0])
            st.metric("Covered Chunks", covered_chunks)
        
        with col3:
            coverage_pct = (covered_chunks / total_chunks) * 100 if total_chunks > 0 else 0
            st.metric("Coverage %", f"{coverage_pct:.1f}%")
        
        with col4:
            avg_questions = sum(chunk_stats.values()) / total_chunks if total_chunks > 0 else 0
            st.metric("Avg Q/Chunk", f"{avg_questions:.1f}")

    @staticmethod
    def show_adaptive_quiz_summary(score: int, total: int, difficulty: str, time_taken: float, wrong_questions: List):
        """Enhanced quiz summary with adaptive insights"""
        
        percentage = (score / total) * 100 if total > 0 else 0
        
        # Performance gauge
        fig_gauge = go.Figure(go.Indicator(
            mode = "gauge+number+delta",
            value = percentage,
            domain = {'x': [0, 1], 'y': [0, 1]},
            title = {'text': "🎯 Quiz Performance"},
            delta = {'reference': 70, 'increasing': {'color': 'green'}},
            gauge = {
                'axis': {'range': [None, 100]},
                'bar': {'color': "darkblue"},
                'steps': [
                    {'range': [0, 50], 'color': "lightgray"},
                    {'range': [50, 75], 'color': "yellow"},
                    {'range': [75, 90], 'color': "orange"},
                    {'range': [90, 100], 'color': "green"}
                ],
                'threshold': {
                    'line': {'color': "red", 'width': 4},
                    'thickness': 0.75,
                    'value': 85
                }
            }
        ))
        
        fig_gauge.update_layout(height=300)
        st.plotly_chart(fig_gauge, use_container_width=True)
        
        # Detailed statistics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Final Score", f"{score}/{total}", f"{percentage:.1f}%")
        
        with col2:
            difficulty_color = {"easy": "🟢", "medium": "🟡", "hard": "🔴"}
            st.metric("Difficulty", f"{difficulty_color.get(difficulty, '⚪')} {difficulty.title()}")
        
        with col3:
            st.metric("Time Taken", f"{time_taken:.1f}s", f"{time_taken/total:.1f}s/Q" if total > 0 else "0s/Q")
        
        with col4:
            wrong_count = len(wrong_questions) if wrong_questions else 0
            st.metric("Wrong Answers", wrong_count, f"{wrong_count/total*100:.1f}%" if total > 0 else "0%")
        
        # Adaptive feedback
        st.subheader("🧠 AI Performance Analysis")
        
        if percentage >= 90:
            st.success("🏆 **Outstanding Performance!** You've mastered this material at the " + difficulty + " level. Consider challenging yourself with harder difficulty!")
        elif percentage >= 75:
            st.info("👍 **Good Understanding!** You grasp most concepts well. Review the missed questions to achieve mastery.")
        elif percentage >= 60:
            st.warning("📚 **Moderate Performance.** Focus on reviewing the material, especially the areas you missed. Consider easier difficulty for practice.")
        else:
            st.error("🔄 **Needs Review.** This material requires more study. Try the retry round and consider switching to easy mode for better learning.")
        
        # Difficulty progression suggestion
        if percentage >= 85 and difficulty == "easy":
            st.info("💡 **Progression Tip:** You're ready for medium difficulty!")
        elif percentage >= 85 and difficulty == "medium":
            st.info("🚀 **Challenge Yourself:** Try hard difficulty next time!")

def save_sensai_quiz_attempt(session_id: str, file_hash: str, filename: str, difficulty: str, 
                           score: int, total: int, chunks_covered: List[int], 
                           wrong_questions: List, time_taken: float):
    """Save quiz attempt with all sensai features"""
    try:
        conn = sqlite3.connect('sensai_quiz.db')
        c = conn.cursor()
        
        c.execute('''INSERT INTO quiz_attempts 
                     (session_id, user_id, file_hash, filename, difficulty, score, total_questions, 
                      timestamp, chunks_covered, wrong_questions, time_taken)
                     VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                  (session_id, "user", file_hash, filename, difficulty, score, total, 
                   datetime.now().isoformat(), json.dumps(chunks_covered), 
                   json.dumps(wrong_questions), time_taken))
        
        # Mark session as completed
        c.execute('''UPDATE chat_sessions SET completed_at = ? WHERE session_id = ?''',
                  (datetime.now().isoformat(), session_id))
        
        conn.commit()
        conn.close()
        log_debug("INFO", f"Quiz attempt saved: {score}/{total}")
        
    except Exception as e:
        log_debug("ERROR", "Failed to save quiz attempt", str(e))

def get_sensai_quiz_history():
    """Get comprehensive quiz history"""
    try:
        conn = sqlite3.connect('sensai_quiz.db')
        df = pd.read_sql_query('''
            SELECT session_id, filename, difficulty, score, total_questions, 
                   (score * 100.0 / total_questions) as percentage,
                   time_taken, timestamp
            FROM quiz_attempts 
            ORDER BY timestamp DESC
        ''', conn)
        conn.close()
        return df
    except Exception as e:
        log_debug("ERROR", "Failed to get quiz history", str(e))
        return pd.DataFrame()

def show_debug_panel():
    """Show debug information panel"""
    if st.sidebar.checkbox("🐛 Show Debug Info"):
        with st.sidebar.expander("Debug Logs", expanded=False):
            try:
                conn = sqlite3.connect('sensai_quiz.db')
                logs_df = pd.read_sql_query('''
                    SELECT timestamp, level, message 
                    FROM debug_logs 
                    ORDER BY timestamp DESC 
                    LIMIT 20
                ''', conn)
                conn.close()
                
                if not logs_df.empty:
                    for _, log in logs_df.iterrows():
                        level_emoji = {"ERROR": "❌", "WARNING": "⚠️", "INFO": "ℹ️"}.get(log['level'], "📝")
                        st.text(f"{level_emoji} {log['message']}")
                else:
                    st.text("No debug logs found")
            except Exception as e:
                st.error(f"Debug panel error: {str(e)}")

def test_api_connection():
    """Test API connection and show status"""
    if st.session_state.openai_api_key:
        try:
            client = initialize_openai_client(st.session_state.openai_api_key)
            if client:
                # Test with a simple request
                response = client.chat.completions.create(
                    model="openai/gpt-4o-mini",
                    messages=[{"role": "user", "content": "Test connection. Respond with 'OK'."}],
                    max_tokens=5,
                    temperature=0
                )
                if response:
                    st.sidebar.success("✅ API Connection: Active")
                    return True
        except Exception as e:
            st.sidebar.error(f"❌ API Connection: Failed - {str(e)}")
            log_debug("ERROR", "API connection test failed", str(e))
            return False
    return False

def main():
    st.title("🧠 Sensai AI - Doc→Assessment Bot")
    st.markdown("### 🚀 Transform Documents into Intelligent Conversational Quizzes")
    
    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Sensai Configuration")
        
        # API Key input
        if not st.session_state.openai_api_key:
            api_key = st.text_input("🔑 Hackathon API Key", type="password", 
                                  help="Use the API key provided for the hackathon")
            if api_key:
                st.session_state.openai_api_key = api_key
                st.success("✅ API Key configured!")
                st.rerun()
        else:
            st.success("✅ API Key loaded!")
            # Show configuration details
            st.info("🔧 Using hackathon configuration:\n- Base URL: agent.dev.hyperverge.org\n- Model: openai/gpt-4o-mini")
            
            # Test API connection
            if st.button("🔗 Test API Connection"):
                test_api_connection()
        
        # Difficulty selector with adaptive hints
        st.markdown("**🎚️ Quiz Difficulty**")
        difficulty = st.selectbox(
            "Choose difficulty level:",
            ["easy", "medium", "hard"],
            index=1,
            help="Easy: Basic recall, Medium: Comprehension, Hard: Critical analysis"
        )
        
        # Difficulty description
        difficulty_desc = {
            "easy": "🟢 **Basic Understanding** - Definitions and recall",
            "medium": "🟡 **Comprehension** - Concepts and applications", 
            "hard": "🔴 **Critical Analysis** - Complex reasoning and synthesis"
        }
        st.markdown(difficulty_desc[difficulty])
        
        # Quiz settings
        st.markdown("**⚙️ Quiz Settings**")
        max_questions = st.slider("Max Questions per Quiz", 5, 20, 10)
        show_progress = st.checkbox("Show Progress Bar", True)
        
        # Debug panel
        show_debug_panel()
        
        # Navigation
        st.markdown("---")
        if st.button("📊 View Quiz History", use_container_width=True):
            st.session_state.show_history = True
        
        if st.button("🔄 Reset Session", use_container_width=True):
            # Reset all session state except API key
            keys_to_keep = ['openai_api_key']
            keys_to_delete = [k for k in st.session_state.keys() if k not in keys_to_keep]
            for key in keys_to_delete:
                del st.session_state[key]
            st.rerun()
    
    # Initialize session state for sensai features
    if 'sensai_state' not in st.session_state:
        st.session_state.sensai_state = 'upload'
    if 'questions' not in st.session_state:
        st.session_state.questions = []
    if 'current_question' not in st.session_state:
        st.session_state.current_question = 0
    if 'score' not in st.session_state:
        st.session_state.score = 0
    if 'wrong_questions' not in st.session_state:
        st.session_state.wrong_questions = []
    if 'session_id' not in st.session_state:
        st.session_state.session_id = None
    if 'start_time' not in st.session_state:
        st.session_state.start_time = time.time()
    if 'show_history' not in st.session_state:
        st.session_state.show_history = False
    if 'retry_mode' not in st.session_state:
        st.session_state.retry_mode = False
    
    # Show quiz history
    if st.session_state.show_history:
        st.header("📊 Sensai Quiz History & Analytics")
        
        history_df = get_sensai_quiz_history()
        if not history_df.empty:
            # Summary stats
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Total Quizzes", len(history_df))
            with col2:
                avg_score = history_df['percentage'].mean()
                st.metric("Average Score", f"{avg_score:.1f}%")
            with col3:
                best_score = history_df['percentage'].max()
                st.metric("Best Score", f"{best_score:.1f}%")
            with col4:
                total_time = history_df['time_taken'].sum()
                st.metric("Total Study Time", f"{total_time:.0f}s")
            
            # Performance trend
            fig_trend = px.line(history_df, x='timestamp', y='percentage', 
                              title='📈 Performance Trend Over Time',
                              color='difficulty')
            st.plotly_chart(fig_trend, use_container_width=True)
            
            # Detailed history table
            st.subheader("📋 Detailed History")
            st.dataframe(history_df, use_container_width=True)
        else:
            st.info("📝 No quiz history found. Take your first quiz to see analytics here!")
        
        if st.button("⬅️ Back to Quiz", use_container_width=True):
            st.session_state.show_history = False
            st.rerun()
        return
    
    # Main Sensai workflow
    if st.session_state.sensai_state == 'upload':
        # Upload interface
        st.header("📤 Upload Your Document")
        st.markdown("Upload any **PDF** or **DOCX** file to generate an intelligent quiz with AI-powered questions and citations.")
        
        uploaded_file = st.file_uploader(
            "Choose your document:",
            type=['pdf', 'docx'],
            help="Supported: PDF and DOCX files up to 10 pages for optimal performance"
        )
        
        if uploaded_file is not None:
            # File validation and preview
            file_size = len(uploaded_file.read())
            uploaded_file.seek(0)  # Reset file pointer
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.info(f"📄 **File:** {uploaded_file.name}")
            with col2:
                st.info(f"📏 **Size:** {file_size/1024:.1f} KB")
            with col3:
                file_type = "PDF" if uploaded_file.type == "application/pdf" else "DOCX"
                st.info(f"📋 **Type:** {file_type}")
            
            # Token estimation
            rough_chars = file_size if uploaded_file.type != "application/pdf" else file_size * 2
            estimated_tokens = rough_chars / 4  # Rough token estimation
            
            if estimated_tokens > 50000:
                st.warning(f"⚠️ Large document detected (~{estimated_tokens/1000:.0f}K tokens). Processing will be limited to prevent API overuse.")
            
            # Process button
            if st.button("🚀 Generate Sensai Quiz", use_container_width=True, type="primary"):
                if not st.session_state.openai_api_key:
                    st.error("🔑 Please provide your hackathon API key in the sidebar!")
                    return
                
                # Processing pipeline with comprehensive error handling
                with st.spinner("🔄 Processing document with Sensai AI..."):
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    try:
                        # Step 1: Extract text
                        status_text.text("📖 Extracting text from document...")
                        progress_bar.progress(20)
                        
                        if uploaded_file.type == "application/pdf":
                            text, page_count, page_content = SensaiDocumentProcessor.extract_text_from_pdf(uploaded_file)
                        else:
                            text, page_count, page_content = SensaiDocumentProcessor.extract_text_from_docx(uploaded_file)
                        
                        if not text or len(text.strip()) < 100:
                            st.error("❌ Could not extract sufficient text from document. Please try a different file.")
                            log_debug("ERROR", "Insufficient text extracted", f"Text length: {len(text)}")
                            return
                        
                        # Step 2: Chunk text
                        status_text.text("✂️ Chunking text with citations...")
                        progress_bar.progress(40)
                        
                        chunks = SensaiDocumentProcessor.chunk_text_with_citations(text, page_content)
                        if not chunks:
                            st.error("❌ Could not process document into chunks. Please try a different file.")
                            log_debug("ERROR", "No chunks created")
                            return
                        
                        st.session_state.chunks = chunks
                        st.session_state.page_content = page_content
                        
                        # Step 3: Generate questions
                        status_text.text("🧠 Generating intelligent questions with AI...")
                        progress_bar.progress(60)
                        
                        all_questions = []
                        chunks_to_process = min(len(chunks), max_questions // 2)  # Ensure we don't exceed max questions
                        
                        for i, chunk in enumerate(chunks[:chunks_to_process]):
                            status_text.text(f"🎯 Processing chunk {i+1}/{chunks_to_process}...")
                            
                            questions = SensaiQuizGenerator.generate_mcqs_with_citations(chunk, difficulty)
                            if questions:
                                all_questions.extend(questions)
                                log_debug("INFO", f"Generated {len(questions)} questions from chunk {i}")
                            else:
                                log_debug("WARNING", f"No questions generated from chunk {i}")
                            
                            # Update progress
                            progress = 60 + (i + 1) / chunks_to_process * 30
                            progress_bar.progress(int(progress))
                            status_text.text(f"🎯 Generated {len(all_questions)} questions so far...")
                            
                            # Limit total questions
                            if len(all_questions) >= max_questions:
                                break
                        
                        # Step 4: Setup quiz session
                        status_text.text("🎮 Setting up quiz session...")
                        progress_bar.progress(90)
                        
                        if all_questions:
                            st.session_state.questions = all_questions[:max_questions]
                            session_id, question_order = SensaiQuizEngine.create_quiz_session(
                                st.session_state.questions, difficulty
                            )
                            
                            if session_id:
                                st.session_state.session_id = session_id
                                st.session_state.question_order = question_order
                                st.session_state.file_hash = hashlib.md5(uploaded_file.read()).hexdigest()
                                st.session_state.filename = uploaded_file.name
                                st.session_state.difficulty = difficulty
                                st.session_state.start_time = time.time()
                                st.session_state.sensai_state = 'quiz'
                                
                                progress_bar.progress(100)
                                status_text.text("✅ Quiz ready!")
                                
                                # Success message
                                st.success(f"""
                                🎉 **Sensai Quiz Generated Successfully!**
                                - **Questions Created:** {len(st.session_state.questions)}
                                - **Chunks Processed:** {len(chunks)}
                                - **Pages Analyzed:** {page_count}
                                - **Difficulty:** {difficulty.title()}
                                """)
                                
                                time.sleep(2)
                                st.rerun()
                            else:
                                st.error("❌ Failed to create quiz session. Please try again.")
                        else:
                            st.error("❌ No questions could be generated. This could be due to:")
                            st.markdown("""
                            - **API Connection Issues**: Check your API key and connection
                            - **Document Content**: The document may not have enough readable text
                            - **Processing Errors**: Try with a simpler document first
                            
                            **Troubleshooting Tips:**
                            - Ensure your API key is correct
                            - Try a document with clear, structured text
                            - Check the debug panel for more details
                            """)
                            log_debug("ERROR", "No questions generated", f"Chunks: {len(chunks)}, Text length: {len(text)}")
                    
                    except Exception as e:
                        st.error(f"❌ Processing failed: {str(e)}")
                        log_debug("ERROR", "Document processing failed", str(e))
    
    elif st.session_state.sensai_state == 'quiz':
        # Quiz interface
        if st.session_state.retry_mode:
            st.header("🔄 Retry Round - Master Your Missed Questions")
            st.info(f"📚 Reviewing {len(st.session_state.questions)} questions you missed earlier.")
        else:
            st.header("🎯 Sensai Conversational Quiz")
            st.markdown("Answer questions in this chat-like interface. Get instant feedback and hints!")
        
        # Progress tracking
        if show_progress:
            progress = st.session_state.current_question / len(st.session_state.questions)
            st.progress(progress)
        
        # Quiz stats
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Current Score", st.session_state.score)
        with col2:
            st.metric("Progress", f"{st.session_state.current_question}/{len(st.session_state.questions)}")
        with col3:
            accuracy = (st.session_state.score / max(1, st.session_state.current_question)) * 100
            st.metric("Accuracy", f"{accuracy:.0f}%")
        with col4:
            elapsed = time.time() - st.session_state.start_time
            st.metric("Time", f"{elapsed:.0f}s")
        
        # Quiz content
        if st.session_state.current_question < len(st.session_state.questions):
            current_q = st.session_state.questions[st.session_state.current_question]
            
            # Display question in conversational format
            selected_answer = SensaiChatBot.display_question_conversational(
                current_q, 
                st.session_state.current_question + 1, 
                len(st.session_state.questions)
            )
            
            # Process answer
            if selected_answer:
                is_correct = selected_answer == current_q['correct']
                
                # Show immediate feedback
                SensaiChatBot.show_conversational_feedback(is_correct, current_q, selected_answer)
                
                # Update score and track wrong answers
                if is_correct:
                    st.session_state.score += 1
                else:
                    st.session_state.wrong_questions.append(current_q)
                
                # Update session in database
                SensaiQuizEngine.update_session(
                    st.session_state.session_id,
                    st.session_state.current_question + 1,
                    st.session_state.score,
                    st.session_state.wrong_questions
                )
                
                # Auto-advance with delay for reading
                if st.button("➡️ Continue to Next Question", use_container_width=True, type="primary"):
                    st.session_state.current_question += 1
                    if hasattr(st.session_state, 'show_hint'):
                        del st.session_state.show_hint
                    st.rerun()
        else:
            # Quiz completed
            st.session_state.sensai_state = 'completed'
            st.rerun()
    
    elif st.session_state.sensai_state == 'completed':
        # Quiz completion and results
        if st.session_state.retry_mode:
            st.header("🎊 Retry Round Completed!")
            st.markdown("Great job working through those challenging questions again!")
        else:
            st.header("🎉 Sensai Quiz Completed!")
            st.markdown("Congratulations on completing your intelligent quiz!")
        
        # Calculate final metrics
        time_taken = time.time() - st.session_state.start_time
        percentage = (st.session_state.score / len(st.session_state.questions)) * 100
        
        # Adaptive quiz summary with AI insights
        SensaiVisualizer.show_adaptive_quiz_summary(
            st.session_state.score,
            len(st.session_state.questions), 
            st.session_state.difficulty,
            time_taken,
            st.session_state.wrong_questions
        )
        
        # Save attempt to database
        if not st.session_state.retry_mode:
            chunks_covered = list(set([q.get('chunk_id', 0) for q in st.session_state.questions]))
            save_sensai_quiz_attempt(
                st.session_state.session_id,
                st.session_state.file_hash,
                st.session_state.filename,
                st.session_state.difficulty,
                st.session_state.score,
                len(st.session_state.questions),
                chunks_covered,
                st.session_state.wrong_questions,
                time_taken
            )
        
        # Coverage analysis
        if hasattr(st.session_state, 'chunks') and not st.session_state.retry_mode:
            st.header("📊 Document Coverage Analysis")
            st.markdown("See which parts of your document were covered in the quiz:")
            SensaiVisualizer.show_coverage_heatmap(st.session_state.chunks, st.session_state.questions)
        
        # Retry wrong questions feature
        if st.session_state.wrong_questions and not st.session_state.retry_mode:
            st.header("🔄 Master Your Mistakes")
            
            col1, col2 = st.columns(2)
            with col1:
                st.error(f"📝 You missed **{len(st.session_state.wrong_questions)}** questions")
                st.markdown("Reviewing these will help reinforce your learning!")
            
            with col2:
                if st.button("🚀 Start Retry Round", use_container_width=True, type="primary"):
                    # Setup retry session
                    st.session_state.questions = st.session_state.wrong_questions.copy()
                    st.session_state.current_question = 0
                    st.session_state.score = 0
                    st.session_state.retry_mode = True
                    st.session_state.start_time = time.time()
                    st.session_state.sensai_state = 'quiz'
                    
                    # Clear wrong questions for retry
                    st.session_state.wrong_questions = []
                    st.rerun()
        
        # AI Doubt Resolution Feature
        st.header("❓ Ask Your AI Tutor")
        st.markdown("Have questions about the material? Get AI-powered explanations!")
        
        user_doubt = st.text_area(
            "What would you like to understand better?",
            placeholder="e.g., Can you explain the concept from question 3 in more detail?",
            height=100
        )
        
        if st.button("🤖 Get AI Explanation", use_container_width=True) and user_doubt:
            if st.session_state.openai_api_key:
                with st.spinner("🧠 AI Tutor is thinking..."):
                    try:
                        # Use document chunks as context
                        context = ""
                        if hasattr(st.session_state, 'chunks'):
                            context = "\n\n".join([chunk['text'][:1000] for chunk in st.session_state.chunks[:3]])
                        
                        # Initialize OpenAI client with hackathon configuration
                        client = initialize_openai_client(st.session_state.openai_api_key)
                        
                        if client:
                            response = client.chat.completions.create(
                                model="openai/gpt-4o-mini",
                                messages=[
                                    {
                                        "role": "system", 
                                        "content": "You are Sensai, an intelligent tutoring AI. Help students understand concepts from their study material. Be encouraging, clear, and provide examples when helpful."
                                    },
                                    {
                                        "role": "user", 
                                        "content": f"Study Material Context:\n{context}\n\nStudent Question: {user_doubt}\n\nProvide a helpful, encouraging explanation:"
                                    }
                                ],
                                max_tokens=600,
                                temperature=0.7
                            )
                            
                            with st.chat_message("assistant", avatar="🤖"):
                                st.markdown("### 🎓 Sensai AI Tutor Response:")
                                st.markdown(response.choices[0].message.content)
                                
                                # Suggest related study topics
                                st.markdown("---")
                                st.markdown("💡 **Study Suggestions:**")
                                st.markdown("- Review the source material sections highlighted in the quiz")
                                st.markdown("- Try the quiz again on a different difficulty level") 
                                st.markdown("- Focus on the concepts from questions you missed")
                        else:
                            st.error("❌ Could not initialize AI Tutor. Please check your API key.")
                        
                    except Exception as e:
                        st.error(f"❌ AI Tutor error: {str(e)}")
                        log_debug("ERROR", "AI Tutor failed", str(e))
            else:
                st.warning("🔑 Hackathon API key required for AI tutoring features.")
        
        # Action buttons
        st.markdown("---")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("🔄 Take New Quiz", use_container_width=True, type="primary"):
                # Reset for new quiz
                keys_to_keep = ['openai_api_key']
                keys_to_delete = [k for k in st.session_state.keys() if k not in keys_to_keep]
                for key in keys_to_delete:
                    del st.session_state[key]
                st.rerun()
        
        with col2:
            if st.button("📊 View All History", use_container_width=True):
                st.session_state.show_history = True
                st.rerun()
        
        with col3:
            if st.button("📤 Upload New Document", use_container_width=True):
                # Reset to upload state
                st.session_state.sensai_state = 'upload'
                st.session_state.current_question = 0
                st.session_state.score = 0
                st.session_state.questions = []
                st.session_state.wrong_questions = []
                st.rerun()

    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #666; padding: 20px;'>
        <p>🧠 <strong>Sensai AI</strong> - Intelligent Document-to-Quiz Transformation</p>
        <p>🚀 Features: Adaptive Difficulty • Citation-Based Questions • Coverage Analysis • AI Tutoring</p>
        <p>🔧 <strong>Hackathon Configuration:</strong> Using agent.dev.hyperverge.org with GPT-4o-mini</p>
        <p>🐛 <strong>Debug:</strong> Enable debug panel in sidebar for troubleshooting</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
